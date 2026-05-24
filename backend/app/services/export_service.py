"""导出服务:把 Markdown 文本转换成可下载的 Markdown 或 DOCX 字节流。

DOCX 渲染只覆盖项目里实际用到的少量 Markdown 语法,不追求全功能:
- ATX 标题(# / ## / ###)
- 无序列表(- / *)和有序列表(1. / 2.)
- 段落,空行作为段落分隔
- 行内粗体 **xxx** 与斜体 *xxx*

更复杂的 Markdown(表格、代码块、链接等)按段落原样输出,
对当前简历 / 求职信 / 面试准备的内容形态已经足够。
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass
from urllib.parse import quote

from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Pt
from docx.text.paragraph import Paragraph


SUPPORTED_FORMATS = ("markdown", "docx")


@dataclass
class ExportPayload:
    """导出结果三元组,由路由直接拼成 StreamingResponse。"""

    content: bytes
    media_type: str
    filename: str


# ============ Markdown ============


def export_markdown(title: str, body: str) -> ExportPayload:
    """直接把 Markdown 文本编码为下载文件。"""
    return ExportPayload(
        content=body.encode("utf-8"),
        media_type="text/markdown; charset=utf-8",
        filename=_safe_filename(title, "md"),
    )


# ============ DOCX ============


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
_UL_RE = re.compile(r"^[-*]\s+(.+)$")
_OL_RE = re.compile(r"^(\d+)[.\)]\s+(.+)$")
_INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")


def export_docx(title: str, body: str) -> ExportPayload:
    """把 Markdown 文本渲染成 DOCX 字节流。"""
    document: DocxDocument = Document()
    _apply_base_styles(document)

    document.add_heading(title, level=0)
    _render_markdown(document, body)

    buffer = io.BytesIO()
    document.save(buffer)
    return ExportPayload(
        content=buffer.getvalue(),
        media_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        filename=_safe_filename(title, "docx"),
    )


def _apply_base_styles(document: DocxDocument) -> None:
    normal = document.styles["Normal"]
    if normal.font:
        normal.font.size = Pt(11)


def _render_markdown(document: DocxDocument, body: str) -> None:
    """按行解析 Markdown 主体并写入文档。"""
    for raw_line in body.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            # 空行单纯作为段落分隔,DOCX 自带段间距,不需要额外空段
            continue

        heading_match = _HEADING_RE.match(line)
        if heading_match:
            level = min(len(heading_match.group(1)), 6)
            document.add_heading(heading_match.group(2).strip(), level=level)
            continue

        ul_match = _UL_RE.match(line)
        if ul_match:
            paragraph = document.add_paragraph(style="List Bullet")
            _write_inline(paragraph, ul_match.group(1).strip())
            continue

        ol_match = _OL_RE.match(line)
        if ol_match:
            paragraph = document.add_paragraph(style="List Number")
            _write_inline(paragraph, ol_match.group(2).strip())
            continue

        paragraph = document.add_paragraph()
        _write_inline(paragraph, line)


def _write_inline(paragraph: Paragraph, text: str) -> None:
    """把行内文本按 **粗体** / *斜体* 切片写入 run,其它原样输出。"""
    for token in _INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) > 4:
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)


# ============ 文件名清理 ============


_INVALID_FILENAME_CHARS = re.compile(r'[\\/:*?"<>|\r\n\t]+')


def _safe_filename(title: str, extension: str) -> str:
    """把任意 title 压成跨平台合法的文件名(保留中文)。"""
    cleaned = _INVALID_FILENAME_CHARS.sub(" ", title).strip()
    if not cleaned:
        cleaned = "export"
    # 限长,留出扩展名空间
    return f"{cleaned[:80]}.{extension}"


def build_content_disposition(filename: str) -> str:
    """RFC 6266: 非 ASCII 文件名要用 filename* + UTF-8 百分号编码。

    保留 ASCII fallback,旧浏览器拿不到 filename* 时仍能存盘。
    """
    encoded = quote(filename, safe="")
    return f"attachment; filename=\"export\"; filename*=UTF-8''{encoded}"
