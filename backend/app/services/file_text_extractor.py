"""通用文件到纯文本的抽取模块。

当前支持 PDF / DOCX / TXT / Markdown 上传，被两类业务复用：
- 简历上传（POST /api/v1/resumes/upload）：7'a
- 知识库文档上传（POST /api/v1/knowledge/.../documents）：7'c1

范围刻意保持小而清晰：只做文本抽取。扫描件 PDF 会得到空文本或接近空文本，
OCR 兜底留到后续迭代；加密或损坏文件会向调用方返回可直接展示给用户的错误。

本模块也刻意不绑定框架，不引 FastAPI / SQLAlchemy。调用方捕获
:class:`FileExtractionError` 后，再自行决定映射为 HTTP 400、聊天回复或其他业务结果。
"""

from __future__ import annotations

import io
from dataclasses import dataclass


# 5 MB 上限已经明显高于常见纯文本简历（约 100 KB）和大多数招聘页面 PDF。
# 如果真实用户频繁撞到这个限制，后续可以按部署环境改成 env 配置。
MAX_FILE_BYTES = 5 * 1024 * 1024

# 归一化后的文本如果短于这个阈值，通常意味着 OCR-only PDF、空文件，
# 或“看似成功但实际没有抽到内容”的解析失败。这里尽早暴露问题，
# 避免把垃圾输入继续送进 LLM 解析或 RAG 切片。
MIN_EXTRACTED_CHARS = 30


class FileExtractionError(Exception):
    """上传内容无法可靠转换为纯文本时抛出。

    ``user_message`` 会被 API 原样返回，因此保持短小、明确、可行动，
    前端无需再做二次翻译即可展示。
    """

    def __init__(self, user_message: str) -> None:
        super().__init__(user_message)
        self.user_message = user_message


@dataclass(frozen=True)
class ExtractedFile:
    text: str
    source_type: str  # "pdf" | "docx" | "markdown" | "text"


def extract_text_from_upload(
    filename: str,
    content_type: str | None,
    payload: bytes,
) -> ExtractedFile:
    """优先按文件扩展名分发，必要时再回落到 content_type。

    文件扩展名通常比浏览器提供的 MIME type 更可靠，下载文件尤其容易被标成
    ``application/octet-stream``。因此 content_type 只作为无扩展名文件的兜底判断。
    """
    if len(payload) > MAX_FILE_BYTES:
        size_mb = MAX_FILE_BYTES // (1024 * 1024)
        raise FileExtractionError(
            f"文件超过 {size_mb} MB 上限，请压缩或截取后再上传。",
        )

    if not payload:
        raise FileExtractionError("上传的文件是空的。")

    lower = filename.lower().strip()
    if lower.endswith(".pdf"):
        text = _extract_pdf(payload)
        source = "pdf"
    elif lower.endswith(".docx"):
        text = _extract_docx(payload)
        source = "docx"
    elif lower.endswith((".md", ".markdown")):
        text = _decode_text(payload)
        source = "markdown"
    elif lower.endswith(".txt"):
        text = _decode_text(payload)
        source = "text"
    elif content_type in ("application/pdf",):
        text = _extract_pdf(payload)
        source = "pdf"
    elif content_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ):
        text = _extract_docx(payload)
        source = "docx"
    elif content_type and content_type.startswith("text/"):
        text = _decode_text(payload)
        source = "text"
    else:
        raise FileExtractionError(
            "暂不支持这种文件类型，请上传 PDF、DOCX、TXT 或 Markdown。",
        )

    normalised = _normalise_whitespace(text)
    if not normalised:
        if source in {"text", "markdown"}:
            raise FileExtractionError("上传的文本文件没有内容。")
        raise FileExtractionError(
            "文件里抽不到可读文本（可能是扫描件或图片型 PDF）。"
            "请改上传可复制文本的版本，或先转存为 DOCX/TXT。",
        )
    if source in {"pdf", "docx"} and len(normalised) < MIN_EXTRACTED_CHARS:
        raise FileExtractionError(
            "文件里抽不到可读文本（可能是扫描件或图片型 PDF）。"
            "请改上传可复制文本的版本，或先转存为 DOCX/TXT。",
        )
    return ExtractedFile(text=normalised, source_type=source)


def _extract_pdf(payload: bytes) -> str:
    """使用 pdfplumber 从 PDF 中抽取文本。

    这里采用懒加载，避免后端其他路径承担导入成本。pdfplumber 会连带导入
    pdfminer.six 和 pillow，首次导入通常约几十毫秒。
    """
    import pdfplumber  # noqa: PLC0415 — lazy import for cold-start cost

    try:
        with pdfplumber.open(io.BytesIO(payload)) as pdf:
            chunks = [page.extract_text() or "" for page in pdf.pages]
    except Exception as exc:  # noqa: BLE001 — pdfplumber surfaces many error types
        raise FileExtractionError(
            "PDF 解析失败，可能文件已损坏或加密；请上传其他格式。",
        ) from exc

    return "\n\n".join(chunk.strip() for chunk in chunks if chunk and chunk.strip())


def _extract_docx(payload: bytes) -> str:
    """使用 python-docx 从 DOCX 中抽取文本。

    python-docx 会按文档顺序返回段落；这里用换行拼接，让下游的 LLM 解析和
    RAG 切片能看到接近正文结构的文本。
    """
    from docx import Document  # noqa: PLC0415 — lazy import

    try:
        document = Document(io.BytesIO(payload))
    except Exception as exc:  # noqa: BLE001 — python-docx raises broad exceptions
        raise FileExtractionError(
            "DOCX 解析失败，可能文件已损坏；请重新导出后再试。",
        ) from exc

    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)

    # python-docx 会把表格和段落分开暴露，因此这里也抽取单元格文本。
    # 很多简历会用侧边栏表格放技能，知识库文档也可能用表格承载结构化信息，
    # 如果忽略表格就会静默丢失关键内容。
    for table in document.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def _decode_text(payload: bytes) -> str:
    """优先按 UTF-8 解码，再温和回落到 GBK / UTF-16。

    现在大多数在线简历和笔记都是 UTF-8，但中文 Windows 环境导出的老文件
    仍可能是 GBK。尝试几个常见编码，可以减少用户因为编码问题被挡住的概率。
    """
    for encoding in ("utf-8", "gbk", "utf-16"):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise FileExtractionError("文件编码不支持，请保存为 UTF-8 后重试。")


def _normalise_whitespace(text: str) -> str:
    """裁剪每行行尾空白，并压缩连续空行。

    Token 预算很宝贵，PDF 尤其容易在段落之间插入大量空行。压缩这些空行不损失信息，
    还能显著缩短下游 LLM prompt。
    """
    lines = [line.rstrip() for line in text.splitlines()]
    out: list[str] = []
    blank_pending = False
    for line in lines:
        if line.strip():
            if blank_pending and out:
                out.append("")
            out.append(line)
            blank_pending = False
        else:
            blank_pending = True
    return "\n".join(out).strip()
